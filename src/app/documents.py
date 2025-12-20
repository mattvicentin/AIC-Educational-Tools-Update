#!/usr/bin/env python3
"""
documents.py
Purpose: Document generation from chat conversations
Status: [ACTIVE]
Created: 2025-09-13
Author: writeian

Handles generating structured documents (notes, outlines) from chat discussions
for educational scaffolding without academic dishonesty concerns.
"""

from flask import (
    Blueprint,
    request,
    redirect,
    url_for,
    flash,
    current_app,
    send_file,
    make_response,
)
from datetime import datetime
import io
import tempfile
import os
from typing import Any, List
from src.app import db
from src.models import Chat, Message
from src.app.access_control import get_current_user, require_chat_access

documents = Blueprint("documents", __name__)


@documents.route("/chat/<int:chat_id>/generate", methods=["POST"])
@require_chat_access
def generate_from_chat(chat_id: int) -> Any:
    """Generate a structured document from chat conversation."""
    try:
        # Debug logging for 400 errors
        current_app.logger.info(f"Document generation request for chat {chat_id}")
        current_app.logger.info(f"Form data: {dict(request.form)}")
        current_app.logger.info(f"Request method: {request.method}")
        chat_obj = Chat.query.get_or_404(chat_id)
        user = get_current_user()
        
        # Get all messages for this chat
        messages = Message.query.filter_by(chat_id=chat_obj.id).order_by(Message.timestamp).all()
        
        # Check minimum message threshold
        if len(messages) < 5:
            flash(f"Need at least 5 messages to generate a document. Current: {len(messages)} messages.", "warning")
            return redirect(url_for("chat.view_chat", chat_id=chat_obj.id))
        
        # Get document type and format from request
        doc_type = request.form.get("doc_type", "notes")
        format_type = request.form.get("format", "text")
        
        # Generate document content based on chat messages
        document_content = generate_document_content(messages, chat_obj, doc_type)
        
        # Create downloadable file based on format
        if format_type == "docx":
            return create_docx_download(document_content, chat_obj, doc_type)
        else:
            # Default to plain text/markdown
            return create_text_download(document_content, chat_obj, doc_type)
            
    except Exception as e:
        current_app.logger.error(f"Error generating document for chat {chat_id}: {e}")
        flash("Failed to generate document. Please try again.", "error")
        return redirect(url_for("chat.view_chat", chat_id=chat_obj.id))


@documents.route("/chat/<int:chat_id>/export-raw", methods=["POST"])
@require_chat_access
def export_raw_chat(chat_id: int) -> Any:
    """Export raw chat conversation with timestamps and usernames."""
    try:
        # Debug logging for 400 errors
        current_app.logger.info(f"Raw chat export request for chat {chat_id}")
        current_app.logger.info(f"Form data: {dict(request.form)}")
        current_app.logger.info(f"Request method: {request.method}")
        chat_obj = Chat.query.get_or_404(chat_id)
        user = get_current_user()
        
        # Get all messages for this chat
        messages = Message.query.filter_by(chat_id=chat_obj.id).order_by(Message.timestamp).all()
        
        # Get format from request
        format_type = request.form.get("format", "txt")
        
        # Generate raw chat content
        raw_content = generate_raw_chat_content(messages, chat_obj)
        
        # Create downloadable file based on format
        if format_type == "docx":
            return create_raw_docx_export(raw_content, chat_obj)
        else:
            # Default to plain text
            return create_raw_text_export(raw_content, chat_obj)
            
    except Exception as e:
        current_app.logger.error(f"Error exporting raw chat {chat_id}: {e}")
        flash("Failed to export chat. Please try again.", "error")
        return redirect(url_for("chat.view_chat", chat_id=chat_obj.id))


def generate_document_content(messages: List[Message], chat_obj: Chat, doc_type: str) -> str:
    """Generate structured document content from chat messages."""
    from src.utils.openai_utils import call_anthropic_api
    from src.app.room.utils.room_utils import infer_template_type_from_room
    
    # Prepare chat content for AI analysis
    chat_content = []
    for msg in messages:
        role_prefix = "User" if msg.role == "user" else "AI Assistant"
        if msg.user and msg.role == "user":
            role_prefix = f"{msg.user.display_name}"
        chat_content.append(f"{role_prefix}: {msg.content}")
    
    chat_text = "\n\n".join(chat_content)
    
    # Determine template type for context
    template_type = infer_template_type_from_room(chat_obj.room) or "general"
    
    # Create AI prompt based on document type
    prompts = {
        "notes": f"""
Based on the following chat discussion from a {template_type} learning session, create comprehensive summary notes that capture what was actually discussed. Focus on summarizing insights, not creating blanks to fill.

Chat Discussion:
{chat_text}

Create summary notes with:
1. **Original Ideas Presented** (highlight human contributions and initial concepts)
2. **Key Insights Developed** (main points that emerged through dialogue)
3. **Technical Solutions Explored** (approaches and methods discussed)
4. **Questions That Emerged** (new questions raised during discussion)
5. **Collaborative Breakthroughs** (ideas that developed through interaction)

Format as a comprehensive summary of what was covered, not as a template to complete. Clearly distinguish between human-originated ideas and AI-provided frameworks.
""",
        "outline": f"""
Based on the following chat discussion from a {template_type} learning session, create a document outline appropriate for the type of work this room is designed to produce.

Chat Discussion:
{chat_text}

Template Context: {template_type}

Create an outline structure that matches the room's purpose:
- For academic-essay: Research paper structure (thesis, literature review, methodology, analysis, conclusion)
- For business-hub: Business document structure (executive summary, analysis, recommendations, implementation)
- For study-group: Study material structure (concepts, applications, practice, review)
- For learning-lab: Technical report structure (overview, methodology, findings, implementation)
- For creative-studio: Creative project structure (concept, development, execution, reflection)
- For writing-workshop: Writing piece structure (planning, drafting, revision, publication)

Include:
- Section headings appropriate for the intended document type
- Framework based on chat insights but structured for the room's goal
- [Student develops] placeholders for sections requiring original work
- Guiding questions specific to the document type

This should be a roadmap for creating the type of document this room is designed to produce.
""",
        "summary": f"""
Based on the following chat discussion, create an organized summary that captures the key insights and structures the thinking. Focus on frameworks and organization, not completed content.

Chat Discussion:
{chat_text}

Create a helpful organizational structure that students can build upon.
"""
    }
    
    prompt = prompts.get(doc_type, prompts["summary"])
    
    try:
        # Generate document content using AI
        if doc_type == "notes":
            system_prompt = "You are an expert educator who creates comprehensive summaries of learning discussions. Focus on capturing what was actually discussed, highlighting human contributions, and organizing insights clearly. Do not create fill-in-the-blank templates."
        else:  # outline
            system_prompt = "You are an expert educator who creates document frameworks appropriate for different types of academic and professional work. Provide structure that matches the intended output type while leaving space for student development."
        
        content, _ = call_anthropic_api(
            [{"role": "user", "content": prompt}],
            system_prompt=system_prompt,
            max_tokens=1500
        )
        return content
    except Exception as e:
        current_app.logger.error(f"Error generating document content: {e}")
        return f"""# Discussion Notes

Generated from chat: {chat_obj.title}
Date: {datetime.now().strftime('%Y-%m-%d')}
Room: {chat_obj.room.name}

## Key Discussion Points

[Content generation failed - please try again]

## Messages in Discussion: {len(messages)}

This document would normally contain organized notes from your chat discussion.
"""


def create_text_download(content: str, chat_obj: Chat, doc_type: str) -> Any:
    """Create a downloadable text file."""
    filename = f"{doc_type}_{chat_obj.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.txt"
    
    # Create response with text content
    response = make_response(content)
    response.headers['Content-Type'] = 'text/plain; charset=utf-8'
    response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    return response


def create_docx_download(content: str, chat_obj: Chat, doc_type: str) -> Any:
    """Create a downloadable Word document."""
    try:
        # Try to import python-docx
        from docx import Document
        from docx.shared import Inches
        
        # Create new document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f"{doc_type.title()}: {chat_obj.title}", 0)
        
        # Add metadata
        doc.add_paragraph(f"Generated from AI Collab chat discussion")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph(f"Room: {chat_obj.room.name}")
        doc.add_paragraph(f"Learning Mode: {chat_obj.mode}")
        doc.add_paragraph("")  # Empty line
        
        # Add content (parse markdown-style formatting)
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('# '):
                # Main heading
                doc.add_heading(line[2:], 1)
            elif line.startswith('## '):
                # Sub heading
                doc.add_heading(line[3:], 2)
            elif line.startswith('### '):
                # Sub-sub heading
                doc.add_heading(line[4:], 3)
            elif line.startswith('- ') or line.startswith('* '):
                # Bullet point
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ')):
                # Numbered list
                doc.add_paragraph(line[3:], style='List Number')
            else:
                # Regular paragraph
                doc.add_paragraph(line)
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        filename = f"{doc_type}_{chat_obj.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        def cleanup_temp_file():
            try:
                os.unlink(temp_file.name)
            except:
                pass
        
        # Create response
        response = send_file(
            temp_file.name,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        # Schedule cleanup
        import atexit
        atexit.register(cleanup_temp_file)
        
        return response
        
    except ImportError:
        # Fallback to text if python-docx not available
        current_app.logger.warning("python-docx not available, falling back to text export")
        return create_text_download(content, chat_obj, doc_type)
    except Exception as e:
        current_app.logger.error(f"Error creating DOCX file: {e}")
        return create_text_download(content, chat_obj, doc_type)


def get_available_document_types(message_count: int, template_type: str = None) -> List[dict]:
    """Get available document types based on message count and room template."""
    types = []
    
    if message_count >= 5:
        types.append({
            "key": "notes",
            "label": "📝 Discussion Notes",
            "description": "Organized notes with key points and questions"
        })
    
    if message_count >= 10:
        types.append({
            "key": "outline", 
            "label": "📋 Document Outline",
            "description": "Structured framework for writing"
        })
        
        # Template-specific options
        if template_type == "academic-essay":
            types.append({
                "key": "essay_framework",
                "label": "📄 Essay Framework", 
                "description": "Essay structure with thesis development"
            })
        elif template_type == "business-hub":
            types.append({
                "key": "business_summary",
                "label": "💼 Business Summary",
                "description": "Executive summary and action items"
            })
    
    return types


def generate_raw_chat_content(messages: List[Message], chat_obj: Chat) -> str:
    """Generate raw chat content with timestamps and usernames."""
    content = f"Chat: {chat_obj.title}\n"
    content += f"Room: {chat_obj.room.name}\n"
    content += f"Exported: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n"
    content += f"Total Messages: {len(messages)}\n"
    content += "=" * 50 + "\n\n"
    
    for msg in messages:
        timestamp = msg.timestamp.strftime("%B %d, %Y at %I:%M %p")
        if msg.role == "user" and msg.user:
            content += f"{msg.user.display_name} ({timestamp}):\n"
        else:
            content += f"AI Assistant ({timestamp}):\n"
        
        content += f"{msg.content}\n\n"
    
    return content


def create_raw_text_export(content: str, chat_obj: Chat) -> Any:
    """Create downloadable raw chat text file with [chat_title]_[date].txt naming."""
    safe_title = chat_obj.title.replace(' ', '_').replace('/', '_').replace('\\', '_')
    filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d')}.txt"
    
    response = make_response(content)
    response.headers['Content-Type'] = 'text/plain; charset=utf-8'
    response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    return response


def create_raw_docx_export(content: str, chat_obj: Chat) -> Any:
    """Create downloadable raw chat Word document with [chat_title]_[date].docx naming."""
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading(f"Chat: {chat_obj.title}", 0)
        doc.add_paragraph(f"Room: {chat_obj.room.name}")
        doc.add_paragraph(f"Exported: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        doc.add_paragraph("")
        
        # Add formatted chat content
        lines = content.split('\n')
        for line in lines:
            if line.strip():
                if '(' in line and ')' in line and (':' in line):
                    # Message header with timestamp - make bold
                    p = doc.add_paragraph()
                    run = p.add_run(line.strip())
                    run.bold = True
                elif line.startswith('='):
                    continue  # Skip separator
                elif line.startswith(('Chat:', 'Room:', 'Exported:', 'Total Messages:')):
                    # Metadata - italic
                    p = doc.add_paragraph()
                    run = p.add_run(line.strip())
                    run.italic = True
                else:
                    # Message content
                    doc.add_paragraph(line.strip())
        
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        safe_title = chat_obj.title.replace(' ', '_').replace('/', '_').replace('\\', '_')
        filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        def cleanup_temp_file():
            try:
                os.unlink(temp_file.name)
            except:
                pass
        
        response = send_file(
            temp_file.name,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        import atexit
        atexit.register(cleanup_temp_file)
        return response
        
    except ImportError:
        return create_raw_text_export(content, chat_obj)
    except Exception as e:
        current_app.logger.error(f"Error creating raw DOCX export: {e}")
        return create_raw_text_export(content, chat_obj)
